from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import UploadedFile, ChatHistory, ChatSession
from .serializers import UploadedFileSerializer, ChatHistorySerializer, ChatSessionSerializer
import os
from rest_framework.generics import ListAPIView
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

import chromadb
import PyPDF2
import pdfplumber
import docx
import json
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from .chat import respond_to_message, draft_from_questions

from langchain_community.chat_models import ChatOllama
from openai import OpenAI
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from django.http import HttpResponse
from langchain_openai import ChatOpenAI
from langchain_core.retrievers import BaseRetriever
from langchain_core.documents import Document
from langchain.tools.retriever import create_retriever_tool

import pandas as pd

os.environ["OPENAI_API_KEY"] = "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"

llm = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0) #ChatOllama(model = 'qwen:0.5b') #base_url="http://ollama:11434"

embeddings = OpenAIEmbeddings(model="text-embedding-3-large") # base_url="http://ollama:11434" this is smallest model, probably not best for embeddings
client = OpenAI() # this is for parsing templates, not used on actual data

# embeddings = OllamaEmbeddings(model="all-minilm")

# Initialize ChromaDB Client
chroma_client = chromadb.PersistentClient(path="./chroma")
try:
    collection = chroma_client.get_collection(name="grant_docs")
except:
    collection = chroma_client.create_collection(name="grant_docs")

# Initialize the text splitter with custom parameters
custom_text_splitter = RecursiveCharacterTextSplitter(
    # Set custom chunk size
    chunk_size = 256,
    chunk_overlap  = 16,
    # Use length of the text as the size measure
    length_function = len,
    )

class ChromaRetriever(BaseRetriever):
    """List of documents to retrieve from."""
    k: int
    """Number of top results to return"""

    # Retrieve similar documents
    def _retrieve_similar_documents_(self, query):
        query_embedding = get_openai_embeddings([query])[0]
        results = collection.query(query_embedding, n_results=self.k)
        print(results)
        return [Document(page_content=chunk) for chunk in results['documents'][0]]

    def _get_relevant_documents(self, query: str):
        return self._retrieve_similar_documents_(query)

class ToolWrapper:
    def __init__(self):
        retriever = ChromaRetriever(k=5)
        # tools
        tool = create_retriever_tool(
            retriever,
            "search_docs",
            "Searches and returns excerpts from the documents that have uploaded by the user.",
            )
        self.tools = [tool]

tools = ToolWrapper().tools

# extracts words from pdf
def extract_text_from_pdf(pdf_path):
    page_texts = []
    with pdfplumber.open(pdf_path) as pdf:
        page_texts = [page.extract_text() for page in pdf.pages]
        return " ".join(page_texts)

# OpenAI Embeddings 
def get_openai_embeddings(texts):
    doc_result = embeddings.embed_documents(texts)
    return doc_result

# Function to read content from PDF files
def read_pdf(file_path):
    content = ""
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            content += page.extract_text()
    return content

# Function to read content from DOCX files
def read_docx(file_path):
    doc = docx.Document(file_path)
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text
    return content

# Function to chunk text using langchain's text_splitter
def chunk_text(sample):
    texts = custom_text_splitter.create_documents([sample])
    return texts

def extract_questions(file_path):
    if(file_path.endswith('.pdf')):
        text = extract_text_from_pdf(file_path)
    elif(file_path.endswith('.docx')):
        text = read_docx(file_path)
    # Make a completion request referencing the uploaded file
    extraction_prompt = """ Given the following PDF text, \n 
    BEGINNING OF PDF:  \n 
    ********************************************************* \n
    {} \n
    ********************************************************* \n
    END OF PDF \n, 
    generate a JSON output which contains a list of 'Question' objects.
    Each question should contain a 1) description of what the question is asking, 
    2) any mention of word count limit or (None if no mention), must be an integer, no text and 
    3) any mention of page limit (None if no mention), must be an integer, no text. 
    Return a JSON object for all questions which require a essay or short-answer response. You should return a 
    list of objects ONLY for those questions which require an essay/short-answer response. You may re-word the question to make it more verbose, 
    it should be clear what the question is asking so an LLM could answer it correctly.
    Not for those which require factual details or few word responses. 
    """.format(text)
    example = """\nYour return should be a JSON object for example a document with 1 question might produce the following output: \n
    {'questions':[\n
        {"description": "Describe your organization or project goals", "word_limit": 1000, "page_limit": 3},\n
    ]}"""
    extraction_prompt += example
    completion = client.chat.completions.create(
    model="gpt-4-1106-preview",
    messages=[
        {"role": "system", "content": "You are a helpful assistant that can read PDFs and extract the relevant requested information in JSON. \
         You must follow the users instructions without adding any unwanted elements or keys to the final json object."},
        {"role": "user", "content": extraction_prompt}
    ],
    response_format={"type": "json_object"}
    )
    questions = json.loads(completion.choices[0].message.content)
    return questions

class ChromaManager():
    def __init__(self, collection):
        self.collection = collection

    # Function to delete from chroma db given id
    def delete_from_chromadb(self, x):
        # Retrieve all entries from the collection
        all_entries = self.collection.get()  # Assuming collection.get() retrieves all entries
        ids = all_entries['ids']
        to_delete = [id for id in ids if id.split("_")[0].strip() in x]
        self.collection.delete(ids=to_delete)
    
    # Function to store documents in ChromaDB
    def store_document_in_chromadb(self, documents):
        embeddings = get_openai_embeddings([doc['content'] for doc in documents])
        ids = [doc['id'] for doc in documents]
        contents = [doc['content'] for doc in documents]

        self.collection.upsert(
            ids=ids,
            embeddings=embeddings,
            documents=contents,
        )

    # Function to parse and store files in ChromaDB
    def parse_and_store_files(self, file_paths, ids):
        documents = []
        for idx, file_path in enumerate(file_paths):
            if file_path.endswith(".pdf"):
                content = read_pdf(file_path)
            elif file_path.endswith(".docx"):
                content = read_docx(file_path)
            else:
                continue
            chunks = chunk_text(content)
            for jdx, chunk in enumerate(chunks):
                doc_id = str(ids[idx])+"_"+str(jdx)
                documents.append({"id": doc_id, "content": chunk.page_content})
        
        self.store_document_in_chromadb(documents)

chroma_manager = ChromaManager(collection)

class FileListView(ListAPIView):
    queryset = UploadedFile.objects.all()
    serializer_class = UploadedFileSerializer

class FileUploadView(APIView):
    def post(self, request, is_grantapp, *args, **kwargs):
        file = request.FILES['file'] 

        # Save the file using default storage and get the full path
        file_name = default_storage.save("uploads/" + file.name, ContentFile(file.read()))
        full_file_path = default_storage.path(file_name)

        # Ensure the file exists
        if not default_storage.exists(file_name):
            return Response({"error": "File not saved correctly"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Create an UploadedFile instance with the file path
        uploaded_file = UploadedFile(filename=file.name, file=file_name)
        uploaded_file.save()

        serializer = UploadedFileSerializer(uploaded_file)

        if is_grantapp:
            questions = extract_questions(full_file_path)
            draft = draft_from_questions(llm, questions, tools)
            
            # Create the PDF
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = []

            for question, answer in draft.items():
                elements.append(Paragraph(f"<b>{question}:</b>", styles['Heading2']))
                elements.append(Paragraph(answer, styles['BodyText']))
                elements.append(Spacer(1, 12))  # Add space between questions

            doc.build(elements)

            buffer.seek(0)

            # Return the PDF in the response
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="grant_application.pdf"'
            uploaded_file.delete()
            default_storage.delete(full_file_path)
            return response
        else:
            chroma_manager.parse_and_store_files([full_file_path], [uploaded_file.id])
            return Response(serializer.data, status=status.HTTP_201_CREATED)

class FileDeleteView(APIView):
    def delete(self, request, pk, *args, **kwargs):
        file = get_object_or_404(UploadedFile, pk=pk)
        file_path = file.file.path
        default_storage.delete(file_path)

        if os.path.exists(file_path):
            os.remove(file_path)
        try:
            chroma_manager.delete_from_chromadb([str(pk)])
        except:
            pass

        file.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

# delete a chat session
class ChatSessionDeleteView(APIView):
    def delete(self, request, pk, *args, **kwargs):
        session = get_object_or_404(ChatSession, pk=pk)
        session.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

# gets chat history for a given session
class ChatHistoryView(APIView):
    def get(self, request, *args, **kwargs):
        session_id = request.query_params.get("session_id")
        try:
            chat_session = ChatSession.objects.get(id=session_id)
            chat_history = ChatHistory.objects.filter(session=chat_session)
            serializer = ChatHistorySerializer(chat_history, many=True)
            return Response(serializer.data)
        except ChatSession.DoesNotExist:
            return Response({"error": "Chat session not found"}, status=404)


class LlmModelView(APIView):
    def post(self, request, *args, **kwargs):
        # Save message in chat history
        message = request.data.get("body")
        session_id = request.data.get("session_id")
        chat_session = ChatSession.objects.get(id=session_id)

        chat_history = ChatHistory(user="user", message=message, session=chat_session)
        chat_history.save()

        response = respond_to_message(llm, message, tools, chat_session)

        # Save response in chat history
        chat_history = ChatHistory(user="assistant", message=response, session=chat_session)
        chat_history.save()
        return Response({"response" : response})

class ChatSessionView(ListAPIView):
    queryset = ChatSession.objects.all()
    serializer_class = ChatSessionSerializer

# add new Chat Session
class ChatSessionCreateView(APIView):
    def post(self, request, *args, **kwargs):
        name = request.data.get("body")
        session = ChatSession(name=name)
        session.save()
        return Response(status=status.HTTP_201_CREATED)